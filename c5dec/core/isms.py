from copy import deepcopy
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.text.run import Run
from docx.oxml.text.run import CT_R
import re
import pandas as pd
import os.path
from datetime import datetime
from time import time
import json
import pathlib
import numpy as np
import pandas as pd
import re
import os 
import csv
import warnings
import c5dec.settings as c5settings
warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')

class WordTagProcessor:
    """
    The model class for the MS Word tag processor.
    """

    def __init__(self) -> None:
        self.params_are_set = False
        self.ignore_missing_tag_mapping = False
        self.default_link = "https://google.com"

        self.regex_text = "(.*?)(#S(?:-\w+)+)(.*?)"
        self.tag_regex = re.compile(r"{}".format(self.regex_text))

    def set_params(self, doc_path, csv_path, regex_text, ignore_missing_tag=False):
        self.doc_path = doc_path
        self.csv_path = csv_path
        self.regex_text = regex_text
        self.tag_regex = re.compile(r"{}".format(regex_text))
        self.keep_style = False
        self.ignore_missing_tag_mapping = ignore_missing_tag
        self.output_name = ""

        self.doc = docx.Document(doc_path)
        self.paras = self.doc.paragraphs

        self.tag_dictionary = self.read_csv_into_dict(filepath=csv_path)

        if self.output_name == "":
            self.output_name = 'processed-'+os.path.basename(doc_path)+'.docx'

        self.params_are_set = True

        if self.tag_dictionary is None:
            self.params_are_set = False

    @staticmethod
    def create_run_element():
        # Create a w:r element
        run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        run.append(rPr)

        return run

    @staticmethod
    def create_text_element(text):
        # Create a w:t text element
        new_text = docx.oxml.shared.OxmlElement('w:t')
        tPr = docx.oxml.shared.OxmlElement('w:tPr')
        new_text.append(tPr)
        new_text.text = text

        return new_text

    @staticmethod
    def create_hyperlink_element(p, link):
        # Create a hyperlink element
        r_id = p.part.relate_to(
            link, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        return hyperlink

    @staticmethod
    def create_hyperlink_run(hyperlink_element, link_text, para, run):

        new_run_element = para._element._new_r()
        run._element.addnext(new_run_element)
        new_run = Run(new_run_element, run._parent)
        new_run.text = link_text

        hyperlink_element.append(new_run_element)
        run._r.append(para.add_run(" ")._element)
        run._r.append(hyperlink_element)

        new_run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        new_run.font.underline = True

    def read_csv_into_dict(self, filepath):
        try:
            tag_link_dict = pd.read_csv(
                filepath, header=None, index_col=0).squeeze("columns").to_dict()
        except Exception:
            return None

        return tag_link_dict

    @staticmethod
    def delete_para(p):
        p = p._element
        p.getparent().remove(p)
        p._p = p._element = None

    def extend_para(self, p, text, tag, link, skip_tag=False):
        """
        Adds a hyperlink to the tag in paragraph p, pointing to link.
        """
        # Create a run
        new_text_run = self.create_run_element()

        # Create a new run for the original text in front of the tag
        new_text = self.create_text_element(text)
        new_text_run.append(new_text)

        # Add an extra run in between to add a space character and then add the text run to the paragraph
        r = p.add_run(new_text_run.text, new_text_run.style)
        p.add_run(" ")

        if skip_tag:
            # Create a run
            new_text_run = self.create_run_element()

            # Create a new run for the original text in front of the tag
            new_text = self.create_text_element(tag)
            new_text_run.append(new_text)

            # Add an extra run in between to add a space character and then add the text run to the paragraph
            p.add_run(new_text_run.text, new_text_run.style)
            p.add_run(" ")
        else:
            # Create a hyperlink element
            hyperlink = self.create_hyperlink_element(p, link)
            self.create_hyperlink_run(hyperlink, tag, p, r)

    def rebuild_para(self, p, text1, tag, text2, link):
        """
        Adds a hyperlink to the tag in paragraph p, pointing to link.
        """
        if text1 != "":
            # Create a run
            new_text_run = self.create_run_element()

            # Create a new run for the original text in front of the tag
            new_text = self.create_text_element(text1)
            new_text_run.append(new_text)

            # Add an extra run in between to add a space character and then add the text run to the paragraph
            r = p.add_run()
            r._r.append(new_text_run)
            p.add_run(" ")

        # Create a hyperlink element
        hyperlink = self.create_hyperlink_element(p, link)

        # Create a run (w:r) element and a new properties (w:rPr) element
        new_run = self.create_run_element()

        # Join all xml elements and update the hyperlink text
        new_run.text = tag
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink to it
        r = p.add_run()
        r._r.append(hyperlink)

        # Workaround for when hyperlink style missing; delete if using a template that has a hyperlink style
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        # Create a w:r element and a new w:rPr element
        new_text_run = self.create_run_element()

        # Create a new run for the original text in front of the tag
        new_text = self.create_text_element(text2)
        new_text_run.append(new_text)

        # Add an extra run in between to add a space character and then add the text run to the paragraph
        p.add_run(" ")
        r = p.add_run()
        r._r.append(new_text_run)

    def process_paras_for_tags(self):
        """
        For each regex match, i.e., a compatible hash tag pattern, remove the
        hash tag and replace it with a hyperlink (in place).
        This algorithm builds on a run-based logic (see Microsoft OXML specs for runs)
        by going over runs per paragraph, retrieving the first pattern hit, adding a 
        hyperlink based on the found tag to the previous run and removing the tag from 
        the run in which it was found.
        This implementation preserves the original paragraph style info.
        """
        paras = self.paras
        tag_link_dict = self.tag_dictionary
        tag_regex = self.tag_regex
        total_matches = 0
        for p in paras:
            for i, run in enumerate(p.runs):
                match = tag_regex.search(run.text)
                if match:
                    total_matches += 1
                    previous_run = None
                    try:
                        previous_run = p.runs[i-1]
                    except IndexError:
                        previous_run = run

                    tag = match.group(2).strip()
                    link = tag_link_dict.get(tag)
                    if link is None:
                        if self.ignore_missing_tag_mapping:
                            continue
                        else:
                            link = self.default_link

                    leading_tag = False
                    if run.text.find(tag) == 0:
                        leading_tag = True

                    run.text = run.text.replace(tag, "")

                    hyperlink = self.create_hyperlink_element(p, link)

                    # Handle corner case where tag appears at the very
                    # beginning of a run: add tag to new run; append a
                    # copy of the original run
                    if i == 0 and leading_tag:
                        p_copy = deepcopy(p)
                        p.clear()
                        new_run_element = p._element._new_r()
                        r = p.add_run()
                        r._r.append(new_run_element)
                        previous_run = r

                        self.create_hyperlink_run(
                            hyperlink, tag, p, previous_run)

                        for r in p_copy.runs:
                            new_r = p.add_run()
                            new_r._r.append(r._element)
                    else:
                        self.create_hyperlink_run(
                            hyperlink, tag, p, previous_run)

    def rebuild_paras_for_tags(self):
            """
            Makes hyperlinks based on hash tags and tag to link mapping by
            rebuilding the paragraphs.
            Does not preserve style information (more reliable solution). 
            """
            paras = self.paras
            tag_link_dict = self.tag_dictionary
            tag_regex = re.compile(r"(.*?)(#S(?:-\w+)+)(.*)")
            for p in paras:

                t = str(p.text)
                match = tag_regex.search(t)
                match_found = False
                match_count = 0

                while match:
                    match_count += 1
                    match_found = True
                    if match_count == 1:
                        new_p = p.insert_paragraph_before()
                        p.clear()
                    text_after_tag = ""

                    text_before_tag = match.group(1).strip()
                    tag = match.group(2).strip()
                    text_after_tag = match.group(3).strip()

                    t = text_after_tag
                    
                    link = tag_link_dict.get(tag)
                    skip_tag = False
                    if link is None:
                        if self.ignore_missing_tag_mapping:
                            skip_tag = True
                        else:
                            link = self.default_link

                    self.extend_para(new_p, text_before_tag, tag, link, skip_tag=skip_tag)

                    match = tag_regex.search(t)

                if match_found:
                    new_p.add_run(t)
                    self.delete_para(p)

    def convert_tags_to_hyperlinks(self, keep_style=False):
        if not self.params_are_set:
            raise IOError
        if keep_style:
            self.process_paras_for_tags()
        else:
            self.rebuild_paras_for_tags()
        self.doc.save(os.path.join('./', self.output_name))

class DocListAssistant:
    """The model of the doc list assistant.
    
    This is the model part of the MVC strcuture.
    """
    def __init__(self):
        """Set up the key attributes needed for doc list assistance.
        
        Get all authors from the json file in which they are stored.
        """
        self.doclist = list()
        self.doclist_path = ""
        self.doc_scan_path = ""
        self.unlisted_docs = list()

        self.used_filename_column_name = "UsedFilename"

        self.folders = dict()

    def get_unlisted_docs(self):
        self.unlisted_docs = list()
        df = pd.read_excel(self.doclist_path, sheet_name='DocList')
    
        folder = pathlib.Path(self.doc_scan_path)

        indexed_files_map = dict()
        if not self.used_filename_column_name in df:
            self.used_filename_column_name = "Filename"
            return list()
        
        for fn in df[self.used_filename_column_name]:
            indexed_files_map[fn] = True
        content_obj = folder.rglob("*")
        file_list = list(filter(lambda item: item.is_file(), content_obj))
        namepath_list = list()
        scanned_file_list = []
        for item in file_list:
            pathAndFile = str(item).rsplit(os.sep, 1)
            filename = pathAndFile[1]
            scanned_file_list.append(filename)
            namepath_list.append(pathAndFile)

        unindexed_file_count = 0
        for i, f in enumerate(scanned_file_list):
            if f not in indexed_files_map:
                unindexed_file_count += 1
                self.unlisted_docs.append(f+" at: "+str(namepath_list[i][0]))

        return self.unlisted_docs

    def scandir(self):
        """Scan a directory and collect all files and folders in a 
        dictionary.
        """
        folders = dict()
        for dir, subdirs, files in os.walk(self.path):
            dir = dir[len(self.path):]
            listed_files = []
            for i in files:
                path = self.path + dir + "/" + i
                if os.path.getmtime(path) > self.beginning_date:
                    listed_files.append(i)
            folders.update({dir: listed_files})
        self.folders = folders

    def save_to_csv(self, csv_file):
        """Save the activity report to a csv file.
        
        :param csv_file: the filepath to the csv file
        :type csv_file: str
        """
        with open(csv_file, "w", newline="") as csvfile:
            writer = csv.writer(csvfile, delimiter=",")
            for folder, files in self.activity_report.items():
                for file_data in files:
                    entry = []
                    entry.append(folder+"/"+file_data.pop(0))
                    entry.extend(file_data)
                    writer.writerow(entry)

class ActivityReport:
    """The model of the activity report.
    
    This is the model part of the MVC strcuture for the activity report
    function.
    """
    def __init__(self):
        """Set up the key attributes needed to generate an activity
        report.
        
        Get all authors from the json file in which they are stored.
        """
        self.authors = self.get_authors()
        self.path = ""
        self.author = ""
        self.beginning_date = 0

        self.folders = dict()
        self.activity_report = dict()

    def get_authors(self):
        """Load all authors from the corresponding json file.
        
        :return: the authors and their acronyms
        :rtype: dictionary
        """
        with open(c5settings.PERSON_ACRONYMS_FILE_PATH) as f:
            dictionary = json.load(f)
        
        return dictionary

    def set_date(self, months, days, hours):
        """Set the date after which all files should be collected.
        
        :param months: how many months to look back
        :type months: int 
        :param days: how many days to look back
        :type days: int
        :param hours: how many hours to look back
        :type hours: int
        """
        new_date = time()
        new_date -= (hours*3600)
        new_date -= (days*24*3600)
        for i in range(months):
            year = datetime.fromtimestamp(new_date).year
            month = datetime.fromtimestamp(new_date).month-1
            new_date -= (self.get_days_in_month(month, year)*24*3600)
        self.beginning_date = new_date

    def get_days_in_month(self, month, year):
        """Get the amount of days in a specific month.
        
        :param month: the current month as a number (e.g. June = 6)
        :type month: int
        :param year: the current year
        :type year: int

        :return: the number of days in a month
        :rtype: int
        """
        if month in [1,3,5,7,8,10,12]:
            return 31
        elif month == 2:
            if year%4 == 0 and (year%100 != 0 or year%400 == 0):
                return 29
            else:
                return 28
        else:
            return 30

    def scandir(self):
        """Scan a directory and collect all files and folders in a 
        dictionary.
        """
        folders = dict()
        for dir, subdirs, files in os.walk(self.path):
            dir = dir[len(self.path):]
            listed_files = []
            for i in files:
                path = self.path + dir + "/" + i
                if os.path.getmtime(path) > self.beginning_date:
                    listed_files.append(i)
            folders.update({dir: listed_files})
        self.folders = folders
        
    def get_activity_report(self):
        """Create the activity report of a folder."""
        self.scandir()
        folder_activities = dict()
        for folder, files in self.folders.items():
            activities = []
            for file in files:
                filename = file
                user = self.get_user(filename)
                if self.author != 0:
                    if self.author != user:
                        continue
                date_and_event = self.get_date_and_event(folder, file)
                date = date_and_event[0]
                event = date_and_event[1]
                activities.append([filename, user, date, event])
            folder_activities.update({folder: activities})
        self.activity_report = folder_activities
        
        return self.activity_report
        
    def get_user(self, filename):
        """Get the user/ author of a file by its filename.
        
        :param filename: the name of the file
        :type filename: str

        :return: the author of the file
        :rtype: str
        """
        user = filename[filename.rfind("-")+1:filename.rfind(".")]
        match = False
        for name, acronym in self.authors.items():
            if acronym.lower() == user.lower():
                user = name
                match = True
        if not match:
            user = ""
        return user

    def get_date_and_event(self, folder, file):
        """Get the last modification date and last event of a file.
        
        :param folder: the folder in which the file is stored
        :type folder: str
        :param file: the file name
        :type file: str

        :return: the last modification time and the last event
        :rtype: tuple
        """
        path = self.path + folder + "/" + file
        lmtime = os.path.getmtime(path)
        if lmtime == os.path.getctime(path):
            event = "created"
        else:
            event = "edited"
        lmtime = datetime.fromtimestamp(lmtime)
        lmtime = lmtime.strftime("%d/%m/%Y %H:%M")
        return (lmtime, event)

    def save_to_csv(self, csv_file):
        """Save the activity report to a csv file.
        
        :param csv_file: the filepath to the csv file
        :type csv_file: str
        """
        with open(csv_file, "w", newline="") as csvfile:
            writer = csv.writer(csvfile, delimiter=",")
            for folder, files in self.activity_report.items():
                for file_data in files:
                    entry = []
                    entry.append(folder+"/"+file_data.pop(0))
                    entry.extend(file_data)
                    writer.writerow(entry)
